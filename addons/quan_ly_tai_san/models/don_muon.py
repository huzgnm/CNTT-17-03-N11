# -*- coding: utf-8 -*-
from odoo import models, fields, api
import base64
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DonMuon(models.AbstractModel):
    _name = 'don_muon'
    _description = 'Xuất Đơn Mượn (DOCX)'

    def _add_bold_paragraph(self, doc, text, alignment=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        if alignment:
            p.alignment = alignment
        return p

    def generate_muon_tra_docx(self, muon_tra):
        doc = Document()

        self._add_bold_paragraph(doc, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", WD_PARAGRAPH_ALIGNMENT.CENTER)
        self._add_bold_paragraph(doc, "Độc lập – Tự do – Hạnh phúc", WD_PARAGRAPH_ALIGNMENT.CENTER)
        p = doc.add_paragraph("----------------------")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self._add_bold_paragraph(doc, "HỢP ĐỒNG CHO MƯỢN TÀI SẢN", WD_PARAGRAPH_ALIGNMENT.CENTER)

        p = doc.add_paragraph(
            f"Hôm nay, ngày {muon_tra.thoi_gian_muon.day} "
            f"tháng {muon_tra.thoi_gian_muon.month} "
            f"năm {muon_tra.thoi_gian_muon.year}"
        )
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p = doc.add_paragraph("Tại: ..........................................................")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self._add_bold_paragraph(doc, "\nBên cho mượn")
        doc.add_paragraph("Công ty: Công ty TNHH ABC")
        doc.add_paragraph("Địa chỉ: 123 Đường ABC, Quận XYZ, TP. HCM")
        doc.add_paragraph("Người đại diện: Nguyễn Văn A (Chức vụ: Giám đốc)")

        self._add_bold_paragraph(doc, "\nBên đi mượn")
        doc.add_paragraph(f"Họ tên: {muon_tra.ho_va_ten}")
        doc.add_paragraph(f"Phòng ban: {muon_tra.phong_ban_id.ten_phong_ban}")
        doc.add_paragraph(f"Chức vụ: {muon_tra.chuc_vu_id.ten_chuc_vu}")
        doc.add_paragraph(f"Mã nhân viên: {muon_tra.nhan_vien_id.ma_dinh_danh}")

        self._add_bold_paragraph(doc, "\nĐiều 1: Đối tượng của hợp đồng")
        doc.add_paragraph(
            f"- Bên A đồng ý cho bên B mượn tài sản: "
            f"{muon_tra.ten_tai_san} (Mã: {muon_tra.tai_san_id.ma_tai_san})"
        )
        doc.add_paragraph("- Số lượng: 1")

        self._add_bold_paragraph(doc, "Điều 2: Thời hạn của hợp đồng")
        doc.add_paragraph(f"- Ngày mượn: {muon_tra.thoi_gian_muon}")
        doc.add_paragraph(f"- Ngày trả dự kiến: {muon_tra.thoi_gian_tra_du_kien}")

        self._add_bold_paragraph(doc, "Điều 3: Trách nhiệm của các bên")
        doc.add_paragraph("- Bên B có trách nhiệm bảo quản tài sản trong thời gian sử dụng.")
        doc.add_paragraph("- Khi hoàn trả, tài sản phải ở tình trạng tương đương như lúc nhận.")
        doc.add_paragraph("- Nếu tài sản bị hư hỏng hoặc mất mát, Bên B phải chịu trách nhiệm bồi thường.")

        p = doc.add_paragraph(
            "\nBÊN A (Bên cho mượn)                                   BÊN B (Bên đi mượn)"
        )
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = doc.add_paragraph(
            "\n(Ký tên, đóng dấu)                                     (Ký tên)"
        )
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        docx_file = BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        return docx_file.read()

    def export_muon_tra_docx(self, record_id):
        muon_tra = self.env['muon_tra'].browse(record_id)
        docx_content = self.generate_muon_tra_docx(muon_tra)
        attachment = self.env['ir.attachment'].create({
            'name': f'Hop_Dong_Muon_{muon_tra.ma_muon}.docx',
            'type': 'binary',
            'datas': base64.b64encode(docx_content),
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'res_model': 'muon_tra',
            'res_id': muon_tra.id,
        })
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'self',
        }
